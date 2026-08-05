import io
import re
import pdfplumber
import docx
from PIL import Image
import pytesseract

# ═══════════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION (PDF / DOCX / Image / Fallback)
# ═══════════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(file_bytes):
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Failed to read PDF: {e}")
    
    if not text.strip():
        # Try OCR fallback for image-based / scanned PDFs
        try:
            import pypdfium2 as pdfium
            pdf = pdfium.PdfDocument(file_bytes)
            for i in range(len(pdf)):
                page = pdf[i]
                bitmap = page.render(scale=2)
                pil_image = bitmap.to_pil()
                text += pytesseract.image_to_string(pil_image) + "\n"
        except Exception:
            pass
            
    if not text.strip():
        raise ValueError("PDF contains no extractable text and OCR failed.")
    return text

def extract_text_from_docx(file_bytes):
    text = ""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also extract from tables (many resumes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += row_text + "\n"
    except Exception as e:
        raise ValueError(f"Failed to read DOCX (Note: older .doc format is not supported): {e}")
    return text

def extract_text_from_image(file_bytes):
    text = ""
    try:
        image = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(image)
    except Exception as e:
        raise ValueError(f"Failed to read Image (ensure tesseract-ocr is installed): {e}")
    return text

def extract_text(file_bytes, filename):
    lower_name = filename.lower()
    if lower_name.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif lower_name.endswith('.docx') or lower_name.endswith('.doc'):
        return extract_text_from_docx(file_bytes)
    elif lower_name.endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.webp')):
        return extract_text_from_image(file_bytes)
    else:
        try:
            return file_bytes.decode('utf-8')
        except:
            return ""


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION HEADER DETECTION — robust, handles many resume styles
# ═══════════════════════════════════════════════════════════════════════════════

# Each key is a canonical section name that maps to a standard Excel column.
# Each value is a list of regex patterns that match header variations.
# Patterns are matched case-insensitively against stripped lines.
SECTION_PATTERNS = {
    # ── Objective / Summary / Profile ──────────────────────────────────────
    "Area of Interest / Objective": [
        r"^area\s+of\s+interest",
        r"^career\s+objective",
        r"^objective",
        r"^professional\s+summary",
        r"^summary",
        r"^about\s+me",
        r"^profile\s+summary",
        r"^profile",
        r"^personal\s+statement",
        r"^career\s+goals?",
        r"^introduction",
    ],
    # ── Education ──────────────────────────────────────────────────────────
    "Education": [
        r"^education",
        r"^academic\s+(background|qualifications?|details?|record)",
        r"^educational\s+(background|qualifications?|details?)",
        r"^scholastic\s+record",
        r"^academics?$",
        r"^qualifications?$",
    ],
    # ── Experience & Internships ───────────────────────────────────────────
    "Experience & Internships": [
        r"^work\s+experience",
        r"^professional\s+experience",
        r"^employment\s+(history|experience)",
        r"^experience",
        r"^internships?",
        r"^training\s+&?\s*internships?",
        r"^relevant\s+experience",
        r"^industrial\s+(training|experience)",
        r"^career\s+history",
        r"^job\s+experience",
        r"^positions?\s+held",
    ],
    # ── Projects ──────────────────────────────────────────────────────────
    "Projects": [
        r"^projects?$",
        r"^key\s+projects?",
        r"^academic\s+projects?",
        r"^personal\s+projects?",
        r"^notable\s+projects?",
        r"^selected\s+projects?",
        r"^major\s+projects?",
        r"^mini\s+projects?",
        r"^course\s+projects?",
        r"^capstone\s+projects?",
        r"^research\s+projects?",
        r"^projects?\s+&\s+research",
        r"^projects?\s+(undertaken|completed|done)",
    ],
    # ── Awards & Achievements ──────────────────────────────────────────────
    "Awards & Achievements": [
        r"^awards?\s*/?\s*scholarships?\s*/?\s*(academic\s+)?achievements?",
        r"^awards?\s+(&|and)\s+achievements?",
        r"^awards?(\s+&\s+honors?)?",
        r"^Achievements?(\s+&\s+awards?)?",
        r"^honors?\s+(&|and)\s+awards?",
        r"^scholarships?",
        r"^accolades",
        r"^recognitions?",
        r"^competitive\s+programming",
        r"^coding\s+profiles?",
    ],
    # ── Skills ─────────────────────────────────────────────────────────────
    "Skills": [
        r"^technical\s+skills?",
        r"^skills?$",
        r"^core\s+competenc(ies|e)",
        r"^key\s+skills?",
        r"^skill\s+set",
        r"^technologies?\s+(known|used|stack)?",
        r"^tools?\s+(&|and)\s+technologies?",
        r"^programming\s+(skills?|languages?)",
        r"^expertise",
        r"^proficienc(ies|y)",
        r"^software\s+skills?",
        r"^languages?\s+(&|and)\s+tools?",
        r"^tech\s+stack",
        r"^competenc(ies|e)$",
    ],
    # ── Extra Curriculars & Leadership ─────────────────────────────────────
    "Extra Curriculars & Leadership": [
        r"^positions?\s+of\s+responsibility",
        r"^extra[\s-]?curricular",
        r"^co[\s-]?curricular",
        r"^leadership(\s+&\s+activities)?",
        r"^activities",
        r"^volunteer(ing)?\s*(experience|work)?",
        r"^community\s+(service|involvement)",
        r"^certifications?(\s+&\s+courses?)?",
        r"^courses?\s+(&|and)\s+certifications?",
        r"^professional\s+development",
        r"^memberships?",
        r"^affiliations?",
        r"^hobbies?\s*(&\s*interests?)?",
        r"^interests?$",
        r"^publications?",
        r"^research\s+papers?",
        r"^references?$",
        r"^declaration$",
    ],
}

# Compile all patterns for speed
_COMPILED_PATTERNS = {}
for col, patterns in SECTION_PATTERNS.items():
    _COMPILED_PATTERNS[col] = [re.compile(p, re.IGNORECASE) for p in patterns]


def _clean_header_line(line):
    """
    Normalize a line for header matching:
    strip, remove trailing colons/dashes/underscores, collapse whitespace.
    """
    s = line.strip()
    # Remove common decorators: leading/trailing dashes, underscores, colons, stars, equals
    s = re.sub(r'^[\-=_*#:►▸•]+\s*', '', s)
    s = re.sub(r'\s*[\-=_*#:]+$', '', s)
    s = s.strip()
    return s


def _match_section_header(line):
    """
    Check if a line is a section header. Returns the standard column name
    if matched, or None.
    """
    cleaned = _clean_header_line(line)
    if not cleaned or len(cleaned) > 80:
        # Headers are rarely longer than 80 chars
        return None

    for col, patterns in _COMPILED_PATTERNS.items():
        for pat in patterns:
            if pat.search(cleaned):
                return col
    return None


# ═══════════════════════════════════════════════════════════════════════════════
# CONTACT INFO EXTRACTION — handles diverse formats
# ═══════════════════════════════════════════════════════════════════════════════

# Email: standard pattern
_EMAIL_RE = re.compile(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}')

# Phone: international formats, 10-15 digits with optional +/country code
_PHONE_RE = re.compile(
    r'(?:(?:\+|00)\d{1,3}[\s\-.]?)?\(?\d{2,4}\)?[\s\-.]?\d{3,5}[\s\-.]?\d{3,5}'
)

# LinkedIn
_LINKEDIN_RE = re.compile(r'(?:linkedin\.com/in/|linkedin:\s*)([\w\-]+)', re.IGNORECASE)

# GitHub
_GITHUB_RE = re.compile(r'(?:github\.com/|github:\s*)([\w\-]+)', re.IGNORECASE)

# Registration / Roll / Enrollment number
_REG_PATTERNS = [
    re.compile(r'(?:registration|reg|roll|enrollment|enrolment|id)\s*(?:no|number|#)?\.?\s*:?\s*([\w/\-]+)', re.IGNORECASE),
    re.compile(r'\b(\d{8,11}(?:/\d{4})?)\b'),  # Bare 8-11 digit number with optional /year
]


def _extract_phone(text):
    """Extract phone number, validate it has 10+ digits."""
    for m in _PHONE_RE.finditer(text):
        digits = re.sub(r'\D', '', m.group(0))
        if 10 <= len(digits) <= 15:
            return m.group(0).strip()
    return ""


def _extract_name(lines, email):
    """
    Extract name using heuristics:
    1. First non-empty line that isn't an email/phone/URL/header.
    2. Must look like a name (2-5 words, mostly alpha).
    """
    for line in lines[:5]:  # Name is almost always in first 5 lines
        s = line.strip()
        if not s:
            continue
        # Skip lines that are clearly not names
        if '@' in s or 'http' in s.lower() or 'www.' in s.lower():
            continue
        if re.match(r'^\+?\d[\d\s\-().]{7,}$', s):  # Phone number
            continue
        if _match_section_header(s):  # It's a section header
            continue
        if re.match(r'^(resume|curriculum\s+vitae|cv|biodata)$', s, re.IGNORECASE):
            continue
        # A name typically has 2-5 words, mostly alpha
        words = s.split()
        if 1 <= len(words) <= 6:
            alpha_ratio = sum(1 for w in words if re.match(r'^[A-Za-z.\-\']+$', w)) / len(words)
            if alpha_ratio >= 0.7:
                return s
    return lines[0].strip() if lines else ""


def _extract_registration(text):
    """Extract registration/roll/enrollment number."""
    for pat in _REG_PATTERNS:
        m = pat.search(text)
        if m:
            val = m.group(1)
            # Validate: must contain at least some digits
            if re.search(r'\d{4,}', val):
                return val
    return ""


# ═══════════════════════════════════════════════════════════════════════════════
# DATE PATTERN for project title detection
# ═══════════════════════════════════════════════════════════════════════════════

_DATE_RE = re.compile(
    r'(January|February|March|April|May|June|July|August|September|October|November|December|'
    r'Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}',
    re.IGNORECASE
)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN PARSER
# ═══════════════════════════════════════════════════════════════════════════════

def parse_resume(text):
    """
    Parse extracted resume text into structured fields.
    Handles diverse resume layouts, header styles, and formats.
    """
    data = {
        "Name": "",
        "Contact No": "",
        "Email": "",
        "LinkedIn": "",
        "GitHub": "",
        "Registration No": "",
        "Area of Interest / Objective": "",
        "Education": "",
        "Experience & Internships": "",
        "Projects": "",
        "Awards & Achievements": "",
        "Skills": "",
        "Extra Curriculars & Leadership": "",
        "Role": "Uncategorized"
    }

    if not text.strip():
        return data

    lines = [line.strip() for line in text.split('\n') if line.strip()]

    # ── Contact info extraction (from full text) ───────────────────────
    email_match = _EMAIL_RE.search(text)
    if email_match:
        data["Email"] = email_match.group(0)

    data["Contact No"] = _extract_phone(text)

    linkedin_match = _LINKEDIN_RE.search(text)
    if linkedin_match:
        data["LinkedIn"] = f"linkedin.com/in/{linkedin_match.group(1)}"

    github_match = _GITHUB_RE.search(text)
    if github_match:
        data["GitHub"] = f"github.com/{github_match.group(1)}"

    data["Registration No"] = _extract_registration(text)
    data["Name"] = _extract_name(lines, data["Email"])

    # ── Section detection ──────────────────────────────────────────────
    # Find all section headers and their line indices
    section_hits = []  # list of (line_index, standard_column_name)

    for i, line in enumerate(lines):
        col = _match_section_header(line)
        if col:
            # Avoid detecting the name line as a header
            if i == 0 and data["Name"] and line.strip() == data["Name"]:
                continue
            section_hits.append((i, col))

    # Deduplicate: if same column matched twice, keep first occurrence
    seen_cols = set()
    unique_hits = []
    for idx, col in section_hits:
        if col not in seen_cols:
            seen_cols.add(col)
            unique_hits.append((idx, col))
        else:
            # Allow merging (e.g., both "Experience" and "Internships" sections)
            unique_hits.append((idx, col))

    # Sort by line index
    unique_hits.sort(key=lambda x: x[0])

    # ── Extract content between headers ────────────────────────────────
    for i in range(len(unique_hits)):
        line_idx, col = unique_hits[i]
        start = line_idx + 1

        if i + 1 < len(unique_hits):
            end = unique_hits[i + 1][0]
        else:
            end = len(lines)

        content = "\n".join(lines[start:end]).strip()

        if data[col]:
            data[col] += "\n\n" + content
        else:
            data[col] = content

    # ── Fallback: if no sections detected, try to extract what we can ──
    if not unique_hits:
        # Put everything (except first few contact lines) into a general bucket
        body_start = min(5, len(lines))
        full_body = "\n".join(lines[body_start:])
        data["Skills"] = full_body  # Best guess fallback

    # ── Role determination (uses full text before stripping) ───────────
    data["Role"] = determine_role(
        data["Skills"], data["Projects"],
        data["Experience & Internships"],
        data["Area of Interest / Objective"]
    )

    # ── Post-process Projects: keep only titles ────────────────────────
    if data["Projects"]:
        project_lines = data["Projects"].split("\n")
        titles = []
        for line in project_lines:
            stripped = line.strip()
            if not stripped:
                continue
            # Skip bullet-point description lines
            if stripped[0] in ('•', '–', '-', '▪', '▸', '*', '►', '○', '●'):
                continue
            # A title line has a pipe separator OR a date pattern
            if "|" in stripped or _DATE_RE.search(stripped):
                titles.append(stripped)
        # If the filter removed everything, keep original (different format)
        if titles:
            data["Projects"] = "\n".join(titles)

    return data


# ═══════════════════════════════════════════════════════════════════════════════
# ROLE CLASSIFICATION — keyword-based with weighted scoring
# ═══════════════════════════════════════════════════════════════════════════════

# Each role has keywords with weights. Highest total score wins.
ROLE_KEYWORDS = {
    "Full Stack Developer": {
        "keywords": ["react", "next.js", "angular", "vue", "frontend", "node.js",
                     "express", "django", "backend", "full stack", "fullstack",
                     "mern", "mean", "rest api", "graphql"],
        "weight": 2,
        "requires_both": (
            ["react", "next.js", "angular", "vue", "frontend", "html", "css"],
            ["node.js", "express", "django", "flask", "fastapi", "backend", "mongodb",
             "postgres", "mysql", "spring boot"]
        )
    },
    "Frontend Developer": {
        "keywords": ["react", "next.js", "angular", "vue", "frontend", "html",
                     "css", "tailwind", "sass", "webpack", "ui/ux", "responsive",
                     "bootstrap", "svelte", "redux"],
        "weight": 1,
    },
    "Backend Developer": {
        "keywords": ["node.js", "express", "django", "flask", "fastapi", "backend",
                     "spring boot", "spring", "microservices", "rest api", "graphql",
                     "mongodb", "postgresql", "mysql", "redis", "rabbitmq", "kafka",
                     "grpc", "golang", "go lang"],
        "weight": 1,
    },
    "Data Scientist / ML Engineer": {
        "keywords": ["machine learning", "deep learning", "artificial intelligence",
                     "ai", "tensorflow", "pytorch", "keras", "data science",
                     "nlp", "natural language", "computer vision", "neural network",
                     "scikit", "sklearn", "pandas", "numpy", "data analysis",
                     "regression", "classification", "clustering", "model training",
                     "feature engineering", "hugging face", "transformers",
                     "llm", "large language model", "reinforcement learning"],
        "weight": 1,
    },
    "Data Engineer": {
        "keywords": ["etl", "data pipeline", "data warehouse", "airflow",
                     "spark", "hadoop", "hive", "data lake", "bigquery",
                     "redshift", "snowflake", "dbt", "data engineering",
                     "data ingestion", "batch processing", "stream processing"],
        "weight": 1,
    },
    "DevOps Engineer": {
        "keywords": ["aws", "azure", "gcp", "docker", "kubernetes", "k8s",
                     "ci/cd", "devops", "jenkins", "terraform", "ansible",
                     "github actions", "gitlab ci", "cloud", "infrastructure",
                     "monitoring", "prometheus", "grafana", "linux admin",
                     "site reliability", "sre"],
        "weight": 1,
    },
    "Mobile App Developer": {
        "keywords": ["android", "ios", "flutter", "react native", "swift",
                     "kotlin", "mobile", "app development", "dart",
                     "xcode", "android studio", "firebase"],
        "weight": 1,
    },
    "Cybersecurity Analyst": {
        "keywords": ["cybersecurity", "penetration testing", "pentest",
                     "ethical hacking", "vulnerability", "soc", "siem",
                     "firewall", "ids", "ips", "malware", "forensics",
                     "encryption", "owasp", "burp suite", "nmap",
                     "wireshark", "security audit", "ctf"],
        "weight": 1,
    },
    "UI/UX Designer": {
        "keywords": ["ui design", "ux design", "ui/ux", "figma", "sketch",
                     "adobe xd", "wireframe", "prototype", "user research",
                     "design thinking", "interaction design", "usability",
                     "information architecture"],
        "weight": 1,
    },
    "Embedded Systems / IoT Engineer": {
        "keywords": ["embedded", "iot", "internet of things", "microcontroller",
                     "arduino", "raspberry pi", "rtos", "firmware",
                     "verilog", "vhdl", "fpga", "arm", "sensor",
                     "robotics", "pcb design"],
        "weight": 1,
    },
    "Blockchain Developer": {
        "keywords": ["blockchain", "solidity", "smart contract", "ethereum",
                     "web3", "defi", "nft", "cryptocurrency", "dapp",
                     "hyperledger"],
        "weight": 1,
    },
    "Game Developer": {
        "keywords": ["game development", "unity", "unreal engine", "godot",
                     "game design", "game engine", "opengl", "directx",
                     "3d modeling", "game programming"],
        "weight": 1,
    },
    "Software Development Engineer (SDE)": {
        "keywords": ["data structure", "algorithm", "c++", "java", "python",
                     "software engineer", "sde", "problem solving",
                     "competitive programming", "oops", "oop", "system design",
                     "design patterns", "software development"],
        "weight": 0.5,  # Lower weight — this is the catch-all
    },
}


def determine_role(skills, projects, experience, interests):
    """Classify resume into a role using weighted keyword scoring."""
    combined = f"{skills} {projects} {experience} {interests}".lower()

    scores = {}
    for role, config in ROLE_KEYWORDS.items():
        score = 0
        matched = 0
        for kw in config["keywords"]:
            if kw in combined:
                score += config.get("weight", 1)
                matched += 1

        # Full Stack requires BOTH frontend AND backend keywords
        if role == "Full Stack Developer" and "requires_both" in config:
            front_kws, back_kws = config["requires_both"]
            has_front = any(kw in combined for kw in front_kws)
            has_back = any(kw in combined for kw in back_kws)
            if not (has_front and has_back):
                score = 0  # Disqualify

        if matched > 0:
            scores[role] = score

    if not scores:
        return "Software Development Engineer (SDE)"

    return max(scores, key=scores.get)
